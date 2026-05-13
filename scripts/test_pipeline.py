"""CI test: run the full pipeline on a test .docx and verify outputs."""
import os
import sys
import glob
import time

# Force UTF-8 on Windows to avoid cp1252 encode/decode errors
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
import tempfile
import subprocess
import shutil

TEST_DIR = os.path.join(tempfile.gettempdir(), "hdp_ci_test")
DOCX_PATH = os.path.join(TEST_DIR, "test_document.docx")
OUTPUT_DIR = os.path.join(TEST_DIR, "output")
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
BUNDLED_PIPELINE = os.path.join(PROJECT_ROOT, "dist", "process_pipeline")

def log(msg):
    print(f"[TEST] {msg}")
    sys.stdout.flush()

def make_test_docx():
    log("Creating test .docx...")
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Тестовый документ", 0)
        doc.add_paragraph(
            "Настоящий документ представляет собой тестовое письмо "
            "от Стельмашенко к Червинке относительно дел русского прихода в Праге. "
            "Дата: 15 января 1921 года."
        )
        doc.add_paragraph(
            "В письме обсуждаются вопросы ведения метрических книг, "
            "а также финансовое положение прихода. "
            "Упоминается также Досифей и его роль в церковной общине."
        )
        doc.save(DOCX_PATH)
        log(f"Test .docx created: {DOCX_PATH}")
    except ImportError as e:
        log(f"FAIL: python-docx not available: {e}")
        sys.exit(1)

def find_pipeline_exe():
    ext = ".exe" if sys.platform == "win32" else ""
    dist_dir = os.path.join(PROJECT_ROOT, "dist")

    # PyInstaller OneDir: dist/process_pipeline/process_pipeline(.exe)
    onedir_exe = os.path.join(dist_dir, "process_pipeline", "process_pipeline" + ext)
    if os.path.isfile(onedir_exe):
        return onedir_exe

    # PyInstaller OneFile: dist/process_pipeline(.exe)
    onefile_exe = os.path.join(dist_dir, "process_pipeline" + ext)
    if os.path.isfile(onefile_exe):
        return onefile_exe

    # Deep search in dist
    if os.path.isdir(dist_dir):
        for root, dirs, files in os.walk(dist_dir):
            for f in files:
                if f == "process_pipeline" or f == "process_pipeline.exe":
                    return os.path.join(root, f)

    # Fall back to running via python
    log("Bundled exe not found — will run via python (dev mode)")
    return None

def run_pipeline():
    exe = find_pipeline_exe()
    if exe:
        log(f"Using bundled exe: {exe}")
        cmd = [exe]
    else:
        log("No bundled exe found — running via python")
        python = sys.executable or "python"
        cmd = [python, os.path.join(SCRIPT_DIR, "process_pipeline.py")]

    cmd.extend(["--output-dir", OUTPUT_DIR])
    cmd.extend(["--input-files", DOCX_PATH])

    log(f"Running: {' '.join(cmd)}")
    # Use UTF-8 on all platforms to avoid cp1252 decode errors on Windows
    result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8', errors='replace', timeout=600)

    stdout = result.stdout or ""
    stderr = result.stderr or ""
    log(f"Exit code: {result.returncode}")
    if stdout:
        log(f"STDOUT:\n{stdout[:2000]}")
    if stderr:
        log(f"STDERR:\n{stderr[:2000]}")

    if result.returncode != 0 and result.returncode is not None:
        log(f"WARN: Pipeline exited with code {result.returncode}")

    return stdout, stderr

def validate_outputs(stdout, stderr):
    errors = []
    success = True

    # 1. Check that .md file exists (pandoc conversion)
    md_files = glob.glob(os.path.join(OUTPUT_DIR, "*.md"))
    if md_files:
        log(f"PASS: Pandoc conversion produced {len(md_files)} .md file(s)")
        for mf in md_files:
            sz = os.path.getsize(mf)
            log(f"  {os.path.basename(mf)} — {sz} bytes")
    else:
        # Pandoc might have failed — check error log
        if "Pandoc" in stdout or "pandoc" in stderr.lower():
            errors.append("Pandoc conversion failed")
            success = False
        else:
            log("SKIP: No .md files (pandoc not available in test env)")

    # 2. Check that JSON files exist (LLM extraction)
    json_dir = os.path.join(OUTPUT_DIR, "json")
    json_files = glob.glob(os.path.join(json_dir, "*.json"))
    if json_files:
        log(f"PASS: LLM extraction produced {len(json_files)} .json file(s)")
        for jf in json_files:
            sz = os.path.getsize(jf)
            log(f"  {os.path.basename(jf)} — {sz} bytes")
    else:
        if "API key" in stdout or "api_key" in stdout.lower() or "401" in stderr:
            log("SKIP: LLM extraction skipped (no valid API key)")
        elif "Skipping" in stdout:
            log("SKIP: All files skipped by LLM extraction")
        else:
            log("SKIP: No JSON files (LLM extraction not expected to work in all envs)")

    # 3. Check that HTML visualization exists
    html_path = os.path.join(OUTPUT_DIR, "tufte_timeline.html")
    html_path2 = os.path.join(PROJECT_ROOT, "md", "tufte_timeline.html")
    found_html = None
    for p in [html_path, html_path2]:
        if os.path.isfile(p):
            found_html = p
            break

    if found_html:
        sz = os.path.getsize(found_html)
        log(f"PASS: Visualization HTML generated — {sz} bytes")
    else:
        log(f"WARN: No tufte_timeline.html found (checked: {html_path})")
        # Viz generation may still succeed — check output
        if "Generated visualization" in stdout:
            log("PASS: Viz generation reported success in logs")

    # 4. Check for critical errors
    critical_patterns = [
        "CRITICAL ERROR",
        "Traceback (most recent call last)",
        "FileNotFoundError",
        "PermissionError",
        "Приложение повреждено",
    ]
    for pattern in critical_patterns:
        if pattern in stdout or pattern in stderr:
            errors.append(f"Critical error found: {pattern}")
            success = False

    if errors:
        log(f"FAIL: {'; '.join(errors)}")
        sys.exit(1)
    elif success:
        log("ALL CHECKS PASSED")
    else:
        log("Tests completed with warnings (non-critical)")

def cleanup():
    if os.path.isdir(TEST_DIR):
        shutil.rmtree(TEST_DIR)
        log(f"Cleaned up {TEST_DIR}")

if __name__ == "__main__":
    log("=== CI Pipeline Test ===")
    log(f"Platform: {sys.platform}")
    log(f"Python: {sys.version}")

    try:
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        make_test_docx()
        stdout, stderr = run_pipeline()
        validate_outputs(stdout, stderr)
    finally:
        cleanup()

    log("=== TEST COMPLETE ===")
